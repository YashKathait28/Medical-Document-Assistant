from __future__ import annotations

import io
import uuid
from pathlib import Path
from typing import Dict, List, Tuple

import pandas as pd
import pdfplumber
import pytesseract
from PIL import Image
from docx import Document

from .config import UPLOAD_DIR
from .storage import add_doc
from .utils import chunk_text, safe_filename


def _table_to_text(table: List[List[str]]) -> str:
    rows = []
    for row in table:
        cleaned = [(cell or "").strip() for cell in row]
        rows.append("\t".join(cleaned))
    return "\n".join(rows)


def parse_pdf(path: Path) -> Tuple[str, List[str]]:
    text_parts: List[str] = []
    tables: List[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text:
                text_parts.append(page_text)
            for table in page.extract_tables():
                tables.append(_table_to_text(table))
    return "\n".join(text_parts), tables


def parse_docx(path: Path) -> Tuple[str, List[str]]:
    doc = Document(path)
    text_parts = [para.text for para in doc.paragraphs if para.text]
    tables: List[str] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("\t".join(cells))
        tables.append("\n".join(rows))
    return "\n".join(text_parts), tables


def parse_excel(path: Path) -> Tuple[str, List[str]]:
    dataframes = pd.read_excel(path, sheet_name=None)
    tables = []
    for name, df in dataframes.items():
        tables.append(df.to_csv(sep="\t", index=False))
    return "", tables


def parse_image(path: Path) -> Tuple[str, List[str]]:
    image = Image.open(path)
    text = pytesseract.image_to_string(image)
    return text, []


def ingest_file(
    file_bytes: bytes,
    filename: str,
    source: str,
    source_link: str | None = None,
) -> Dict[str, object]:
    safe_name = safe_filename(filename)
    doc_id = uuid.uuid4().hex
    saved_path = UPLOAD_DIR / f"{doc_id}_{safe_name}"
    with saved_path.open("wb") as handle:
        handle.write(file_bytes)

    extension = saved_path.suffix.lower()
    text = ""
    tables: List[str] = []

    if extension in {".pdf"}:
        text, tables = parse_pdf(saved_path)
    elif extension in {".docx"}:
        text, tables = parse_docx(saved_path)
    elif extension in {".xls", ".xlsx"}:
        text, tables = parse_excel(saved_path)
    elif extension in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
        text, tables = parse_image(saved_path)
    else:
        try:
            text = saved_path.read_text(encoding="utf-8")
        except Exception:
            text = ""

    chunks = chunk_text(text)
    for table in tables:
        if table.strip():
            chunks.append(table)
    doc_meta = {
        "id": doc_id,
        "name": filename,
        "path": str(saved_path),
        "source": source,
        "source_link": source_link,
        "chunks": len(chunks),
        "tables": tables,
    }
    add_doc(doc_meta)
    return {"id": doc_id, "name": filename, "chunks": len(chunks), "chunk_text": chunks}


def build_chunk_payload(doc_id: str, doc_name: str, source_link: str | None, chunks: List[str]) -> Tuple[List[str], List[Dict[str, str]], List[str]]:
    ids = []
    metadatas = []
    for idx, chunk in enumerate(chunks):
        chunk_id = f"{doc_id}_{idx}"
        ids.append(chunk_id)
        metadatas.append(
            {
                "doc_id": doc_id,
                "doc_name": doc_name,
                "chunk_id": chunk_id,
                "source_link": source_link or "",
            }
        )
    return chunks, metadatas, ids
